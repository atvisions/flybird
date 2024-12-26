# doc_generator.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import os

def create_code_style(doc):
    """创建代码块样式"""
    styles = doc.styles
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    paragraph_format = code_style.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    return code_style

def add_code_block(doc, code_style, content):
    """添加代码块"""
    p = doc.add_paragraph()
    p.style = code_style
    p.text = content

def generate_doc():
    doc = Document()
    code_style = create_code_style(doc)
    
    # 1. 标题页
    title = doc.add_heading('档案评分系统技术文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('版本: 1.0.0')
    doc.add_paragraph('日期: 2024-01-10')
    
    # 2. 目录
    doc.add_page_break()
    doc.add_heading('目录', 1)
    
    # 3. 项目概述
    doc.add_page_break()
    doc.add_heading('1. 项目概述', 1)
    doc.add_paragraph('''
档案评分系统是一个基于Django REST framework开发的用户档案质量评估系统。
系统主要功能包括：
• 多维度评分（基础、经验、能力、成就）
• AI内容优化
• 实时评分计算
• 优化建议生成
    ''')
    
    # 4. 系统架构
    doc.add_heading('2. 系统架构', 1)
    
    # 4.1 项目结构
    doc.add_heading('2.1 项目结构', 2)
    add_code_block(doc, code_style, '''
users/
├── backends.py              # 认证后端
├── profile/
│   ├── models/
│   │   ├── score.py        # 评分模型
│   │   └── project.py      # 项目经历模型
│   ├── services/
│   │   └── completeness.py # 评分计算服务
│   ├── views/
│   │   ├── completeness.py # 评分接口
│   │   ├── content_quality.py # 内容质量接口
│   │   └── social.py       # 社交链接接口
│   └── urls.py             # 路由配置
├── urls.py                 # 主路由
└── utils.py               # 工具函数
    ''')
    
    # 5. 核心模型
    doc.add_heading('3. 核心模型', 1)
    
    # 5.1 评分模型
    doc.add_heading('3.1 评分模型 (score.py)', 2)
    with open('users/profile/models/score.py', 'r') as f:
        score_model = f.read()
    add_code_block(doc, code_style, score_model)
    
    # 5.2 评分服务
    doc.add_heading('3.2 评分服务 (completeness.py)', 2)
    with open('users/profile/services/completeness.py', 'r') as f:
        completeness_service = f.read()
    add_code_block(doc, code_style, completeness_service)
    
    # 6. API接口
    doc.add_heading('4. API接口', 1)
    
    # 6.1 评分接口
    doc.add_heading('4.1 评分接口', 2)
    with open('users/profile/views/completeness.py', 'r') as f:
        completeness_view = f.read()
    add_code_block(doc, code_style, completeness_view)
    
    # 6.2 内容质量接口
    doc.add_heading('4.2 内容质量接口', 2)
    with open('users/profile/views/content_quality.py', 'r') as f:
        content_quality_view = f.read()
    add_code_block(doc, code_style, content_quality_view)
    
    # 7. 路由配置
    doc.add_heading('5. 路由配置', 1)
    
    # 7.1 主路由
    doc.add_heading('5.1 主路由 (urls.py)', 2)
    with open('users/urls.py', 'r') as f:
        main_urls = f.read()
    add_code_block(doc, code_style, main_urls)
    
    # 7.2 Profile路由
    doc.add_heading('5.2 Profile路由', 2)
    with open('users/profile/urls.py', 'r') as f:
        profile_urls = f.read()
    add_code_block(doc, code_style, profile_urls)
    
    # 8. 评分算法
    doc.add_heading('6. 评分算法', 1)
    doc.add_paragraph('系统采用多维度加权评分算法：')
    weights = [
        ('基础维度', '40%', '个人基本信息完整度'),
        ('经验维度', '30%', '工作和教育经历'),
        ('能力维度', '20%', '技能和语言能力'),
        ('成就维度', '10%', '证书和作品集'),
        ('内容质量', '额外加分', 'AI优化后的文字质量')
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = '维度'
    header_cells[1].text = '权重'
    header_cells[2].text = '说明'
    
    for dim, weight, desc in weights:
        row_cells = table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = weight
        row_cells[2].text = desc
    
    # 9. 使用示例
    doc.add_heading('7. 使用示例', 1)
    
    # 9.1 获取评分
    doc.add_heading('7.1 获取评分', 2)
    add_code_block(doc, code_style, '''
GET /api/v1/users/profile/completeness/
Authorization: Bearer {token}

Response:
{
    "code": 200,
    "message": "获取成功",
    "data": {
        "total_score": 94.0,
        "total_detail": {
            "basic_dimension": {
                "score": 80.0,
                "weight": 0.4,
                "weighted_score": 32.0
            },
            // ... 其他维度
        }
    }
}
    ''')
    
    # 9.2 优化内容
    doc.add_heading('7.2 优化内容', 2)
    add_code_block(doc, code_style, '''
POST /api/v1/users/profile/content-quality/
Authorization: Bearer {token}

Request:
{
    "field": "personal_summary",
    "content": "优化后的内容...",
    "quality_score": 8.5
}

Response:
{
    "code": 200,
    "message": "内容更新成功",
    "data": {
        "field": "personal_summary",
        "content_quality_score": 8.5
    }
}
    ''')
    
    # 保存文档
    doc.save('档案评分系统技术文档.docx')

if __name__ == '__main__':
    generate_doc()